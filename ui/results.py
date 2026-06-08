import streamlit as st
from io import BytesIO
from docx import Document
from fpdf import FPDF
import re
import html


# ============================================================
# ВСПОМОГАТЕЛЬНЫЕ ФУНКЦИИ (добавлены те, что нужны для HTML)
# ============================================================

def _clean_ai_markdown(text: str) -> str:
    """Удаляет Markdown-разметку из текста."""
    if not text:
        return ""
    text = re.sub(r'^#{1,6}\s+', '', text, flags=re.MULTILINE)
    text = re.sub(r'\*\*(.*?)\*\*', r'\1', text)
    text = re.sub(r'(?<!\n)\*(?!\s)(.*?)(?<!\s)\*', r'\1', text)
    text = re.sub(r'^[\-\*]\s+', '', text, flags=re.MULTILINE)
    text = re.sub(r'\n{3,}', '\n\n', text)
    return text.strip()


def _looks_like_html(value: str) -> bool:
    """Проверяет, содержит ли строка HTML-теги."""
    if not value: return False
    return bool(
        re.search(r"</?(?:p|div|span|strong|em|u|s|h[1-6]|ol|ul|li|blockquote|br)\b", value, flags=re.IGNORECASE))


def _plain_text_to_html(value: str) -> str:
    """Преобразует обычный текст в простые HTML-параграфы."""
    if not value: return "<p><br></p>"
    escaped = html.escape(value)
    paragraphs = re.split(r"\n\s*\n", escaped)
    html_parts = []
    for p in paragraphs:
        p = p.replace("\n", "<br>")
        html_parts.append(f"<p>{p}</p>")
    return "".join(html_parts)


def _build_html_document(content_html: str, title: str) -> str:
    """Создает полноценный HTML-документ со стилями."""
    safe_title = html.escape(title)
    return f"""<!DOCTYPE html>
<html lang="ru">
<head>
    <meta charset="UTF-8">
    <title>{safe_title}</title>
    <style>
        body {{ max-width: 800px; margin: 40px auto; padding: 0 20px; font-family: Arial, sans-serif; line-height: 1.6; color: #333; }}
        h1 {{ border-bottom: 2px solid #eee; padding-bottom: 10px; }}
        p {{ margin-bottom: 1em; }}
    </style>
</head>
<body>
    <h1>{safe_title}</h1>
    <main>{content_html}</main>
</body>
</html>"""


# ============================================================
# ОСНОВНАЯ ФУНКЦИЯ ОТОБРАЖЕНИЯ
# ============================================================

def render_extracted_text(extracted_text: str, file_info: dict):
    """Отображает извлечённый текст и кнопки скачивания"""
    st.subheader("Доступные действия")

    # Очищаем текст один раз
    clean_text = _clean_ai_markdown(extracted_text)

    is_large = file_info.get("size_mb", 0) > 5 or len(clean_text) > 100000
    filename_base = file_info['name'].replace(':', '_')
    base_filename_for_export = f"{filename_base}_text"

    if is_large:
        st.info(f"Текст очень большой. Предпросмотр отключён.")
        # Для больших файлов сразу показываем выбор формата
        _render_additional_formats(clean_text, base_filename_for_export, key_suffix="large")

    else:
        with st.expander("Показать извлеченный текст"):
            st.text(clean_text)

        #  УДАЛИЛИ ОТДЕЛЬНУЮ КНОПКУ TXT (как вы просили)
        # Теперь скачивание только через общий список ниже

        _render_additional_formats(clean_text, base_filename_for_export, key_suffix="small")

    st.write("---")


def _render_additional_formats(content: str, filename_base: str, key_suffix: str):
    """Отрисовывает выбор формата (TXT, PDF, DOCX, HTML)"""
    st.markdown("#### Скачать в формате:")

    export_format = st.selectbox(
        "Выберите формат:",
        ["TXT", "PDF", "DOCX", "HTML"],  # ✅ Добавлен HTML
        label_visibility="collapsed",
        key=f"format_select_{key_suffix}"
    )

    if st.button("Скачать выбранный формат", key=f"btn_download_new_{key_suffix}"):
        if export_format == "TXT":
            download_as_txt(content, f"{filename_base}.txt")
        elif export_format == "PDF":
            download_as_pdf(content, f"{filename_base}.pdf")
        elif export_format == "DOCX":
            download_as_docx(content, f"{filename_base}.docx")
        elif export_format == "HTML":
            # ✅ Логика для HTML
            html_content = content if _looks_like_html(content) else _plain_text_to_html(content)
            html_document = _build_html_document(html_content, "Извлеченный текст")

            st.download_button(
                label="Сохранить HTML",
                data=html_document,
                file_name=f"{filename_base}.html",
                mime="text/html",
                key=f"dl_html_{filename_base}_{key_suffix}"
            )


# --- Функции экспорта (без изменений) ---

def download_as_txt(content: str, filename: str):
    st.download_button(label="️Сохранить TXT", data=content.encode('utf-8'), file_name=filename, mime='text/plain',
                       key=f"dl_txt_{filename}")


def download_as_pdf(content: str, filename: str):
    pdf = FPDF()
    pdf.add_page()
    font_path = "C:/Windows/Fonts/arial.ttf"
    try:
        if not __import__('os').path.exists(font_path): raise FileNotFoundError(font_path)
        pdf.add_font('Arial', '', font_path, uni=True)
        pdf.set_font('Arial', size=12)
        for line in content.split('\n'):
            if len(line) > 100:
                for i in range(0, len(line), 100): pdf.cell(0, 10, txt=line[i:i + 100], ln=True)
            else:
                pdf.cell(0, 10, txt=line, ln=True)
        buffer = BytesIO()
        pdf.output(buffer)
        st.download_button(label="Сохранить PDF", data=buffer.getvalue(), file_name=filename, mime='application/pdf',
                           key=f"dl_pdf_{filename}")
    except Exception as e:
        st.error(f"Ошибка PDF: {e}")


def download_as_docx(content: str, filename: str):
    document = Document()
    document.add_heading('Конспект Vlom', level=1)
    document.add_paragraph(content)
    buffer = BytesIO()
    document.save(buffer)
    st.download_button(label="Сохранить Word", data=buffer.getvalue(), file_name=filename,
                       mime='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
                       key=f"dl_docx_{filename}")