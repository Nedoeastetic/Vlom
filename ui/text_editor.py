import html
import re
from html.parser import HTMLParser
from io import BytesIO
from docx import Document
from fpdf import FPDF

import streamlit as st
from streamlit_quill import st_quill


def _clean_ai_markdown(text: str) -> str:
    """Удаляет Markdown-разметку из ответа ИИ."""
    if not text:
        return ""

    # Убираем заголовки (#, ##)
    text = re.sub(r'^#{1,6}\s+', '', text, flags=re.MULTILINE)
    # Убираем жирный шрифт (**текст**)
    text = re.sub(r'\*\*(.*?)\*\*', r'\1', text)
    # Убираем курсив (*текст*), но аккуратно
    text = re.sub(r'(?<!\n)\*(?!\s)(.*?)(?<!\s)\*', r'\1', text)
    # Убираем маркеры списков (- или *)
    text = re.sub(r'^[\-\*]\s+', '', text, flags=re.MULTILINE)
    # Убираем лишние пустые строки
    text = re.sub(r'\n{3,}', '\n\n', text)

    return text.strip()
# ============================================================
# ВСПОМОГАТЕЛЬНЫЕ ФУНКЦИИ
# ============================================================

class _HTMLTextExtractor(HTMLParser):
    """Преобразует HTML из Quill в обычный текст."""

    BLOCK_TAGS = {
        "p",
        "div",
        "br",
        "li",
        "h1",
        "h2",
        "h3",
        "h4",
        "h5",
        "h6",
        "blockquote",
    }

    def __init__(self) -> None:
        super().__init__()
        self.parts: list[str] = []

    def handle_starttag(
        self,
        tag: str,
        attrs: list[tuple[str, str | None]],
    ) -> None:
        if tag == "br":
            self.parts.append("\n")

        if tag == "li":
            self.parts.append("• ")

    def handle_endtag(self, tag: str) -> None:
        if tag in self.BLOCK_TAGS and tag != "br":
            self.parts.append("\n")

    def handle_data(self, data: str) -> None:
        self.parts.append(data)

    def get_text(self) -> str:
        text = "".join(self.parts)
        text = html.unescape(text)

        # Убираем лишние пробелы перед переносами строк.
        text = re.sub(r"[ \t]+\n", "\n", text)

        # Не допускаем больше двух пустых строк подряд.
        text = re.sub(r"\n{3,}", "\n\n", text)

        return text.strip()


def _init_editor_state() -> None:
    """Инициализирует состояние редактора."""

    defaults = {
        "edit_mode": False,
        "editor_draft": None,
        "editor_original_text": None,
        "result_is_html": False,
        "quill_version": 0,
    }

    for key, value in defaults.items():
        if key not in st.session_state:
            st.session_state[key] = value


def _safe_filename(task_type: str) -> str:
    """Формирует безопасное имя файла."""

    filename = re.sub(
        r"[^\w\-]+",
        "_",
        str(task_type).strip(),
        flags=re.UNICODE,
    )

    filename = filename.strip("_")

    return filename or "result"


def _looks_like_html(value: str) -> bool:
    """Определяет, содержит ли строка HTML-разметку."""

    if not value:
        return False

    return bool(
        re.search(
            r"</?(?:p|div|span|strong|em|u|s|h[1-6]|ol|ul|li|blockquote|br)\b",
            value,
            flags=re.IGNORECASE,
        )
    )


def _plain_text_to_html(value: str) -> str:
    """Преобразует обычный текст в HTML для Quill."""

    if not value:
        return "<p><br></p>"

    escaped = html.escape(value)

    paragraphs = re.split(r"\n\s*\n", escaped)

    html_paragraphs = []

    for paragraph in paragraphs:
        paragraph = paragraph.replace("\n", "<br>")
        html_paragraphs.append(f"<p>{paragraph}</p>")

    return "".join(html_paragraphs)


def _normalize_for_editor(value: str) -> str:
    """Готовит результат для загрузки в Quill."""

    if not value:
        return "<p><br></p>"

    if _looks_like_html(value):
        return value

    return _plain_text_to_html(value)


def _html_to_text(value: str) -> str:
    """Преобразует HTML редактора в TXT."""

    parser = _HTMLTextExtractor()
    parser.feed(value or "")
    parser.close()

    return parser.get_text()


def _is_empty_quill_html(value: str | None) -> bool:
    """Проверяет, пуст ли результат Quill."""

    if not value:
        return True

    plain_text = _html_to_text(value)

    return not plain_text.strip()


def _build_html_document(
    content_html: str,
    task_type: str,
) -> str:
    """Создаёт полноценный HTML-документ для скачивания."""

    safe_title = html.escape(f"Конспект — {task_type}")

    return f"""<!DOCTYPE html>
<html lang="ru">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">

    <title>{safe_title}</title>

    <style>
        body {{
            max-width: 900px;
            margin: 40px auto;
            padding: 0 24px;
            font-family: Arial, sans-serif;
            font-size: 16px;
            line-height: 1.65;
            color: #1f2937;
            background: #ffffff;
            overflow-wrap: anywhere;
        }}

        h1 {{
            margin-bottom: 28px;
        }}

        blockquote {{
            margin: 16px 0;
            padding: 12px 18px;
            border-left: 4px solid #9ca3af;
            background: #f3f4f6;
        }}

        pre {{
            overflow-x: auto;
            padding: 16px;
            border-radius: 8px;
            background: #111827;
            color: #f9fafb;
        }}

        img {{
            max-width: 100%;
            height: auto;
        }}
    </style>
</head>

<body>
    <h1>{safe_title}</h1>

    <main>
        {content_html}
    </main>
</body>
</html>
"""


def _render_result(result: str) -> None:
    """Отображает готовый результат после сохранения."""

    if _looks_like_html(result):
        rendered_content = result
    else:
        rendered_content = _plain_text_to_html(result)

    st.markdown(
        f"""
        <style>
            .vlom-result {{
                padding: 20px;
                border: 1px solid rgba(128, 128, 128, 0.25);
                border-radius: 10px;
                line-height: 1.65;
                overflow-wrap: anywhere;
            }}

            .vlom-result blockquote {{
                margin: 16px 0;
                padding: 10px 16px;
                border-left: 4px solid #9ca3af;
                background: rgba(128, 128, 128, 0.08);
            }}

            .vlom-result img {{
                max-width: 100%;
                height: auto;
            }}
        </style>

        <div class="vlom-result">
            {rendered_content}
        </div>
        """,
        unsafe_allow_html=True,
    )


def _render_save_to_folder(result: str, task_type: str) -> None:
    """Сохраняет текущий конспект в выбранную пользовательскую папку."""

    user_id = st.session_state.get("user_id")

    if not user_id:
        st.info("Войдите в профиль, чтобы сохранять конспекты по папкам.")
        return

    from db.user_manager import create_folder, get_user_folders, save_note

    folders = get_user_folders(user_id)

    st.divider()
    st.markdown("###Сохранить в профиль")

    source_name = (
        st.session_state.get("file_info") or {}
    ).get("name", "без имени")
    default_name = f"{task_type} — {source_name[:40]}"

    note_name = st.text_input(
        "Название конспекта",
        value=default_name,
        max_chars=100,
        key="note_name_input",
    )

    folder_map = {"Без папки": None}
    for folder in folders:
        folder_map[str(folder["name"])] = str(folder["id"])

    folder_options = list(folder_map.keys())

    pending_folder = st.session_state.pop(
        "pending_save_note_folder_select",
        None,
    )
    if pending_folder in folder_options:
        st.session_state["save_note_folder_select"] = pending_folder

    current_folder = st.session_state.get(
        "save_note_folder_select",
        "Без папки",
    )
    if current_folder not in folder_options:
        st.session_state["save_note_folder_select"] = "Без папки"

    selected_folder_name = st.selectbox(
        "Папка",
        options=folder_options,
        key="save_note_folder_select",
        help="Папки создаёт сам пользователь. Готовых папок нет.",
    )

    with st.expander("➕ Создать новую папку перед сохранением"):
        # Поле также очищается только до создания text_input.
        if st.session_state.pop("clear_new_folder_before_save", False):
            st.session_state["new_folder_before_save"] = ""

        new_folder_name = st.text_input(
            "Название новой папки",
            placeholder="Например: Физика",
            max_chars=80,
            key="new_folder_before_save",
        )

        if st.button(
            "Создать папку",
            key="create_folder_before_save_btn",
        ):
            cleaned_folder_name = new_folder_name.strip()

            if not cleaned_folder_name:
                st.warning("Введите название папки.")
            else:
                created = create_folder(user_id, cleaned_folder_name)

                if created:
                    st.session_state[
                        "pending_save_note_folder_select"
                    ] = str(created["name"])
                    st.session_state["clear_new_folder_before_save"] = True
                    st.session_state["folder_success_message"] = (
                        f"Папка «{created['name']}» создана и выбрана."
                    )
                    st.rerun()
                else:
                    st.error(
                        "Не удалось создать папку. Возможно, такое название уже используется."
                    )

    success_message = st.session_state.pop("folder_success_message", None)
    if success_message:
        st.success(success_message)

    if st.button(
        "Сохранить конспект",
        key="save_note_btn",
        type="primary",
        use_container_width=True,
    ):
        name = note_name.strip() or default_name
        folder_id = folder_map[selected_folder_name]

        saved = save_note(
            user_id=user_id,
            note_name=name,
            content=result,
            folder_id=folder_id,
        )

        if saved:
            st.success(f"Конспект «{name}» сохранён.")
        else:
            st.error(
                "Не удалось сохранить конспект. Выполните SQL-миграцию папок в Supabase."
            )


# ============================================================
# ОСНОВНАЯ ФУНКЦИЯ
# ============================================================

def edit_saved_result(task_type: str) -> None:
    """
    Показывает результат ИИ, rich-text редактор и кнопки скачивания.

    В app.py функция должна вызываться так:

        edit_saved_result(task_type)
    """

    _init_editor_state()

    raw_result = st.session_state.get("llm_result")
    # Очищаем текст от звездочек и решеток перед показом
    result = _clean_ai_markdown(raw_result) if raw_result else None

    if not result:
        return

    filename = _safe_filename(task_type)

    st.divider()
    st.success("Результат готов!")

    # ========================================================
    # РЕЖИМ ПРОСМОТРА
    # ========================================================
    if not st.session_state.edit_mode:
        _render_result(result)

        st.divider()

        # Создаем 4 колонки: Редактирование + 3 варианта скачивания
        col_edit, col_download, col_spacer, col_spacer2 = st.columns([1.5, 1, 0.5, 0.5])

        # --- Кнопка редактирования (без изменений) ---
        with col_edit:
            if st.button("️ Редактировать конспект", key="btn_show_edit_panel", use_container_width=True):
                editor_html = _normalize_for_editor(result)
                st.session_state.editor_original_text = result
                st.session_state.editor_draft = editor_html
                st.session_state.result_is_html = _looks_like_html(result)
                st.session_state.edit_mode = True
                st.session_state.quill_version += 1
                st.rerun()

        # --- Блок выбора формата и скачивания ---
        with col_download:
            # Определяем чистый текст и HTML-версию из АКТУАЛЬНОГО session_state
            current_result = st.session_state.get("llm_result") or result
            plain_result = _html_to_text(current_result) if _looks_like_html(current_result) else current_result
            result_html = current_result if _looks_like_html(current_result) else _plain_text_to_html(current_result)

            export_format = st.selectbox(
                "Формат:",
                ["TXT", "PDF", "Word"],
                label_visibility="collapsed",
                key=f"fmt_select_{st.session_state.get('quill_version', 0)}"
            )

            safe_name = f"summary_{task_type.replace(' ', '_')}"

            if export_format == "TXT":
                st.download_button(
                    label="Скачать TXT",
                    data=plain_result.encode('utf-8'),
                    file_name=f"{safe_name}.txt",
                    mime="text/plain",
                    key="dl_txt_dynamic",
                    use_container_width=True,
                )

            elif export_format == "PDF":
                pdf = FPDF()
                pdf.add_page()
                font_path = "C:/Windows/Fonts/arial.ttf"
                try:
                    pdf.add_font('Arial', '', font_path, uni=True)
                    pdf.set_font('Arial', size=12)
                    for line in plain_result.split('\n'):
                        if len(line) > 100:
                            chunks = [line[i:i + 100] for i in range(0, len(line), 100)]
                            for chunk in chunks:
                                pdf.cell(0, 10, txt=chunk, ln=True)
                        else:
                            pdf.cell(0, 10, txt=line, ln=True)
                    buffer = BytesIO()
                    pdf.output(buffer)
                    st.download_button(
                        label="Скачать PDF",
                        data=buffer.getvalue(),
                        file_name=f"{safe_name}.pdf",
                        mime="application/pdf",
                        key="dl_pdf_dynamic",
                        use_container_width=True,
                    )
                except Exception as e:
                    st.error(f"Ошибка PDF: {e}")

            elif export_format == "Word":
                doc = Document()
                doc.add_heading(f'Конспект: {task_type}', level=1)
                doc.add_paragraph(plain_result)
                buffer = BytesIO()
                doc.save(buffer)
                st.download_button(
                    label="⬇ Скачать Word",
                    data=buffer.getvalue(),
                    file_name=f"{safe_name}.docx",
                    mime='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
                    key="dl_docx_dynamic",
                    use_container_width=True,
                )

        _render_save_to_folder(result, task_type)
        return

    # ========================================================
    # РЕЖИМ РЕДАКТИРОВАНИЯ
    # ========================================================

    if st.session_state.editor_draft is None:
        st.session_state.editor_draft = _normalize_for_editor(result)

    if st.session_state.editor_original_text is None:
        st.session_state.editor_original_text = result

    st.markdown("###Редактор конспекта")

    st.caption(
        "Выделите нужный фрагмент прямо в редакторе, затем выберите "
        "жирный шрифт, цвет, размер, заголовок, список или выравнивание "
        "на панели над текстом."
    )

    toolbar = [
        [
            "bold",
            "italic",
            "underline",
            "strike",
        ],
        [
            {"script": "sub"},
            {"script": "super"},
        ],
        [
            {"header": [1, 2, 3, 4, False]},
            {"size": ["small", False, "large", "huge"]},
        ],
        [
            {"font": []},
            {"color": []},
            {"background": []},
        ],
        [
            {"list": "ordered"},
            {"list": "bullet"},
            {"indent": "-1"},
            {"indent": "+1"},
        ],
        [
            {"align": []},
        ],
        [
            "blockquote",
            "code-block",
            "link",
            "clean",
        ],
    ]

    edited_html = st_quill(
        value=st.session_state.editor_draft,
        placeholder="Редактируйте конспект...",
        html=True,
        toolbar=toolbar,
        preserve_whitespace=True,
        key=f"quill_editor_{st.session_state.quill_version}",
    )

    # У некоторых версий компонента первое значение может быть None.
    if edited_html is None:
        edited_html = st.session_state.editor_draft
    else:
        st.session_state.editor_draft = edited_html

    st.divider()

    col_save, col_cancel = st.columns(2)

    with col_save:
        if st.button(
            "Сохранить изменения",
            key="btn_save_edit",
            type="primary",
            use_container_width=True,
        ):
            if _is_empty_quill_html(edited_html):
                st.warning("Конспект не может быть пустым.")
            else:
                st.session_state.llm_result = edited_html
                st.session_state.result_is_html = True
                st.session_state.editor_draft = None
                st.session_state.editor_original_text = None
                st.session_state.edit_mode = False

                st.rerun()

    with col_cancel:
        if st.button(
            "Отменить",
            key="btn_cancel_edit",
            use_container_width=True,
        ):
            original_text = st.session_state.editor_original_text

            if original_text is not None:
                st.session_state.llm_result = original_text

            st.session_state.editor_draft = None
            st.session_state.editor_original_text = None
            st.session_state.edit_mode = False

            # При следующем открытии создастся новый экземпляр Quill.
            st.session_state.quill_version += 1

            st.rerun()

    # ========================================================
    # СКАЧИВАНИЕ ТЕКУЩЕЙ РЕДАКЦИИ
    # ========================================================

    st.divider()
    st.markdown("#### Скачать текущую редакцию")

    current_html = edited_html or st.session_state.editor_draft or ""
    current_text = _html_to_text(current_html)

    col_download_txt, col_download_html = st.columns(2)

    with col_download_txt:
        st.download_button(
            label="Скачать текущий текст в TXT",
            data=current_text,
            file_name=f"summary_{filename}.txt",
            mime="text/plain; charset=utf-8",
            key="download_result_txt_edit",
            use_container_width=True,
        )

